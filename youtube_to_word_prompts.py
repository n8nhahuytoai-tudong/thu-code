#!/usr/bin/env python3
"""
YouTube Scene-by-Scene Analyzer - ONE LINE PER SCENE FOR WORD
Xuất ra file Word, mỗi cảnh = 1 dòng prompt CỰC CHI TIẾT

Features:
- Mỗi scene = 1 prompt duy nhất trên 1 dòng
- Prompt 250-350 words: nhân vật (chiều cao, cân nặng), bối cảnh, âm thanh
- Xuất ra file .docx (Word) và .txt
- Tiêu chuẩn Hollywood blockbuster
"""

import os
import sys
import time
import subprocess
import cv2
import base64
import numpy as np
from pathlib import Path
import json
from openai import OpenAI
from datetime import datetime
import hashlib
from typing import List, Dict, Optional
import shutil

# Import docx - REQUIRED
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    DOCX_AVAILABLE = True
except ImportError:
    print("\n" + "="*70)
    print("ERROR: python-docx chưa được cài đặt!".center(70))
    print("="*70)
    print("\nFile Word (.docx) cần package python-docx để tạo.\n")
    print("Cài đặt ngay:")
    print("  pip install python-docx")
    print("\nHoặc cài tất cả:")
    print("  pip install -r requirements.txt")
    print("\n" + "="*70 + "\n")
    sys.exit(1)

# ========== CONFIGURATION ==========

class Config:
    """Cấu hình"""
    SCENE_THRESHOLD = 30.0
    MIN_SCENE_LENGTH = 15
    MAX_VIDEO_HEIGHT = 1080

    VISION_MODEL = "gpt-4o"
    WHISPER_MODEL = "whisper-1"
    MAX_RETRIES = 5
    RETRY_DELAY = 3

    CACHE_DIR = "cache"
    OUTPUT_DIR = "output_scenes"
    TEMP_FRAMES_DIR = "temp_frames"

    MAX_SCENES_TO_ANALYZE = 999

# ========== UTILITIES ==========

def load_env_file():
    """Load .env file"""
    env_file = Path('.env')
    if env_file.exists():
        with open(env_file, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith('#') and '=' in line:
                    key, value = line.split('=', 1)
                    key, value = key.strip(), value.strip()
                    if value.startswith('"') and value.endswith('"'):
                        value = value[1:-1]
                    if value.startswith("'") and value.endswith("'"):
                        value = value[1:-1]
                    os.environ[key] = value

def ensure_directories():
    """Tạo thư mục"""
    for d in [Config.CACHE_DIR, Config.OUTPUT_DIR, Config.TEMP_FRAMES_DIR]:
        Path(d).mkdir(exist_ok=True)

def cleanup_temp():
    """Xóa temp files"""
    for f in ["temp_video.mp4", "temp_audio.m4a"]:
        if os.path.exists(f):
            try:
                os.remove(f)
            except:
                pass
    if os.path.exists(Config.TEMP_FRAMES_DIR):
        try:
            shutil.rmtree(Config.TEMP_FRAMES_DIR)
            Path(Config.TEMP_FRAMES_DIR).mkdir(exist_ok=True)
        except:
            pass

def print_header(text: str):
    print("\n" + "="*70)
    print(text.center(70))
    print("="*70 + "\n")

def print_progress(msg: str, step: int = None, total: int = None):
    if step and total:
        print(f"[{step}/{total}] {msg}")
    else:
        print(f"• {msg}")

def print_success(msg: str):
    print(f"✓ {msg}")

def print_error(msg: str):
    print(f"✗ {msg}")

# ========== MAIN ANALYZER ==========

class SceneBySceneWordExporter:
    """Phân tích và xuất mỗi scene thành 1 dòng prompt siêu chi tiết"""

    DESCRIPTION_PROMPT = """Analyze this {duration:.1f}s scene and provide a detailed description in Vietnamese.

You are viewing 2 frames: FIRST FRAME and LAST FRAME of this scene.

Provide analysis in the following 7 categories IN VIETNAMESE:

1. **Hành động**: Describe what is happening in the scene, the main action or movement
2. **Nhân vật/Đối tượng**: Describe characters, objects, or creatures visible in the scene
3. **Cảm xúc**: Describe the emotional atmosphere and mood of the scene
4. **Bối cảnh**: Describe the setting, location, and environment
5. **Camera**: Describe camera movement, angle, and shot type
6. **Ánh sáng**: Describe lighting setup, color temperature, and quality
7. **Composition**: Describe frame composition, subject placement, and visual balance

Format your response EXACTLY like this (use numbered list):

1. **Hành động**: [your description here]

2. **Nhân vật/Đối tượng**: [your description here]

3. **Cảm xúc**: [your description here]

4. **Bối cảnh**: [your description here]

5. **Camera**: [your description here]

6. **Ánh sáng**: [your description here]

7. **Composition**: [your description here]"""

    ULTRA_DETAILED_PROMPT = """Analyze this {duration:.1f}s scene and create ONE CONTINUOUS LINE PROMPT (NO LINE BREAKS) for Sora 2, following HOLLYWOOD BLOCKBUSTER STANDARDS.

You are viewing 2 frames: FIRST FRAME and LAST FRAME of this scene.

Create a prompt in ENGLISH, EXTREMELY DETAILED, 250-350 words, MUST include:

📹 CAMERA & CINEMATOGRAPHY (MANDATORY):
- Shot type & size: extreme wide/wide/full shot/medium/close-up/extreme close-up
- Camera movement: static/pan (left/right)/tilt (up/down)/dolly (in/out)/tracking/crane/steadicam/handheld/gimbal - BE SPECIFIC
- Camera angle: eye-level/high-angle/low-angle/dutch-angle/overhead/worm's-eye - WITH DEGREE if tilted
- Lens: focal length in mm (ultra-wide 14-24mm/wide 28-35mm/normal 40-58mm/portrait 85mm/telephoto 100-200mm+)
- Aperture: f-stop (f/1.4 to f/16) and resulting depth of field
- Composition: rule of thirds/golden ratio/centered/symmetrical - specify subject placement
- Aspect ratio: 16:9/2.39:1 anamorphic/2.35:1/1.85:1/IMAX

👥 CHARACTERS - PHYSICAL SPECS (MANDATORY IF PRESENT):
- Count: exact number of people
- Each person MUST include:
  * Gender & age: male/female, approximate age
  * Height: estimate in cm (e.g., 175cm, 185cm, 160cm)
  * Build: slim/average/athletic/muscular/heavy
  * Weight: estimate in kg (e.g., 60kg, 75kg, 90kg)
  * Skin tone: pale ivory/fair/light tan/tan/olive/bronze/brown/deep brown/ebony
  * Hair complete: color (platinum blonde/golden blonde/light brown/dark brown/black/auburn/red/gray/white) + style (straight/wavy/curly/braided/dreadlocks/slicked back/messy) + length (buzz cut/short 3cm/medium 10cm/shoulder-length/long 50cm+)
  * Face: shape (oval/round/square/angular), jawline, cheekbones, eye color, nose, lips
  * Body proportions: head-to-body ratio (1:7/1:8), leg length, torso, shoulder width
  * Costume DETAILED: era (modern/period/futuristic), outfit type, exact colors with hex codes if possible, materials (cotton/leather/silk/denim/metal), fit (tailored/loose/tight), condition (new/worn/damaged)
  * Accessories: jewelry/watches/glasses/hats/bags/weapons - be specific
  * Action & performance: exact movement, body language, facial expression, emotion intensity

🐾 ANIMALS/CREATURES (MANDATORY IF PRESENT):
- Species & breed: be exact (e.g., Golden Retriever, Arabian horse, Bengal cat)
- Size dimensions: height in cm at shoulder, length head-to-tail in cm, weight in kg
- Physical: coat/fur/feathers color (primary + secondary colors), pattern (solid/striped/spotted/mottled), texture (fluffy/sleek/rough)
- Features: ear type, tail type, eye color, distinctive markings
- Body proportions: head:body:legs ratio
- Movement & behavior: gait, speed, action

💡 LIGHTING - HOLLYWOOD SETUP (MANDATORY):
- Overall scheme: 3-point/naturalistic/high-key/low-key/chiaroscuro/Rembrandt/butterfly/split lighting
- Key light: exact position (45° front-left), height (above/eye-level/below), quality (hard/soft), intensity (bright/dim), source type (tungsten/LED/natural)
- Fill light: position, ratio to key (1:2, 1:4, etc.), quality
- Back/rim light: position, creating edge separation or not
- Color temperature: exact Kelvin (2700K warm tungsten/3200K studio tungsten/4000K cool white/5000K daylight/5600K bright sun/6500K overcast/7000K+ shade/blue hour)
- Practicals: visible light sources in scene (lamps/windows/screens/candles/neon signs/fire) with their colors
- Light direction: top/side/front/back - with degrees
- Shadows: sharp/soft/none, shadow direction and length
- Atmosphere: haze/fog/smoke density for volumetric lighting, god rays/light shafts

🎨 COLOR GRADING (MANDATORY):
- Color palette: warm/cool/neutral/complementary/analogous/triadic/monochromatic - specify dominant colors
- LUT style: naturalistic/cinematic teal-orange/bleach bypass/cross-processed/noir/vintage/futuristic
- Saturation: vibrant/normal/muted/desaturated/selective color (which colors popped)
- Contrast: high/medium/low, crushed blacks or lifted blacks
- Tone curve: highlights (blown/retained), midtones, shadows (crushed/visible detail)
- Color science: film emulation (Kodak Vision3/Fuji Eterna) or digital

🏞️ ENVIRONMENT & PRODUCTION DESIGN (MANDATORY):
- Location: interior/exterior, urban/suburban/rural/nature/studio
- Specific setting: bedroom/office/street/forest/beach/warehouse - be detailed
- Architecture: modern/classical/industrial/rustic, materials visible
- Set dressing: furniture, props, decorations - what's in frame
- Background: what's behind subject, depth layers (foreground/mid/background elements)
- Time of day: dawn/morning/noon/afternoon/golden hour/dusk/night/blue hour - with exact lighting match
- Weather: clear/cloudy/overcast/rainy/snowy/foggy/stormy
- Season indicators: summer/autumn/winter/spring visual cues
- VFX elements: CGI (none/minimal/moderate/heavy), what's CGI if any, particle effects, compositing

🎬 ACTION & STORY (MANDATORY):
- What's happening: exact action from start to end of scene
- Character motivation: why are they doing this
- Movement speed: slow motion/real-time/fast/frenetic
- Pacing: contemplative/steady/dynamic/chaotic
- Emotional tone: specific emotion (joy/anger/fear/sadness/surprise/disgust/neutral)
- Mood: atmosphere (tense/calm/exciting/mysterious/romantic/horrific)
- Story beat: setup/inciting incident/rising action/climax/resolution
- Genre: action/drama/comedy/horror/sci-fi/romance/thriller/documentary

🎵 SOUND DESIGN (DESCRIBE AUDIO - MANDATORY):
- Ambient sound: quiet/moderate/loud, type of environment sound (city traffic/nature/wind/rain/ocean/crowd/machinery)
- Diegetic sounds: sounds from visible sources in scene (footsteps/doors/vehicles/voices/music from radio)
- Character sounds: dialogue tone (shouting/whispering/normal), breathing, movement sounds
- Music: if any music implied by scene mood (epic orchestral/tense strings/electronic/rock/silence)
- Sound perspective: close intimate/distant/echoey/muffled/clear
- Audio mood: supporting visual tone or contrasting

🎭 STYLE REFERENCE (MANDATORY):
- Film/director comparison: "in the style of [specific film or director]" (e.g., "like Blade Runner 2049", "Nolan's IMAX cinematography", "Wes Anderson symmetry")
- Production value: student film/indie/mid-budget/studio/AAA blockbuster
- Era influence: 1970s grit/1980s neon/1990s grunge/2000s digital/2020s cinematic

CRITICAL REQUIREMENTS - ALL 8 ELEMENTS MANDATORY:
✅ Write in ENGLISH ONLY
✅ ONE CONTINUOUS LINE - absolutely NO line breaks, NO bullet points, just one flowing paragraph
✅ 250-350 words minimum

MUST INCLUDE ALL 8 SECTIONS (even if "none" for some):

1. CAMERA SPECS (MANDATORY): Must mention focal length (mm), aperture (f-stop), camera movement type, shot size, angle
   Example: "35mm lens at f/2.8, steadicam tracking shot..."

2. CHARACTERS IF PRESENT (MANDATORY IF VISIBLE): height in cm, weight in kg, build type, exact skin tone, hair color+style+length, costume with colors and materials
   Example: "male protagonist 185cm tall 80kg athletic build, tan olive skin, short dark brown hair 3cm messy style, wearing fitted black leather jacket..."

3. ANIMALS IF PRESENT (MANDATORY IF VISIBLE): species, breed, size dimensions (cm/kg), coat colors, patterns
   Example: "Golden Retriever dog, 60cm shoulder height, 30kg weight, golden coat with white chest..."

4. LIGHTING (MANDATORY): setup type, color temperature in Kelvin, quality (hard/soft), practicals visible
   Example: "3-point lighting with hard key from 45° left at 5600K cool daylight, practical window light..."

5. ENVIRONMENT (MANDATORY): location type, set design details, time of day, weather, VFX if any
   Example: "urban alley exterior at night, rain-soaked pavement, neon signs, CGI rain particles..."

6. SOUND DESIGN (MANDATORY): ambient sounds, diegetic sounds from scene, music mood, audio atmosphere
   Example: "ambient city traffic sounds, footsteps on wet pavement, distant sirens, tense electronic music underscoring mood..."

7. ACTION & STORY (MANDATORY): what's happening, movement speed, emotional tone, story purpose
   Example: "character running urgently through alley, fast-paced action, tense fearful emotion, chase sequence..."

8. STYLE REFERENCE (MANDATORY): comparable film/director, production value level
   Example: "in the style of Blade Runner 2049, AAA blockbuster production value..."

If any element is not visible (e.g., no animals), briefly state "no animals present" but ALL 8 categories must be addressed.

Format your response as ONE CONTINUOUS SENTENCE/PARAGRAPH:
PROMPT: [Start with camera specs, flow into characters if present, then animals if present, then lighting, environment, sound, action, and style - all in one continuous flowing paragraph with no line breaks, 250-350 words total]"""

    def __init__(self, api_key: Optional[str] = None):
        load_env_file()

        if api_key:
            os.environ["OPENAI_API_KEY"] = api_key

        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise ValueError("OpenAI API key not found!")

        self.client = OpenAI(api_key=api_key)
        ensure_directories()

        self.video_path: Optional[str] = None
        self.audio_path: Optional[str] = None
        self.youtube_url: Optional[str] = None
        self.video_title: str = ""
        self.video_metadata: Dict = {}
        self.scenes: List[Dict] = []
        self.transcript: Optional[Dict] = None

    # ========== VIDEO DOWNLOAD ==========

    def _get_metadata(self, url: str) -> Dict:
        print_progress("Đang lấy metadata...")
        try:
            result = subprocess.run(
                ["yt-dlp", "--dump-json", url],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0:
                metadata = json.loads(result.stdout)
                self.video_title = metadata.get('title', 'Unknown')
                self.video_metadata = {
                    'title': metadata.get('title'),
                    'duration': metadata.get('duration', 0),
                    'width': metadata.get('width'),
                    'height': metadata.get('height'),
                    'fps': metadata.get('fps'),
                }
                print_success(f"Video: {self.video_title}")
                return self.video_metadata
        except Exception as e:
            print_error(f"Lỗi metadata: {e}")

        self.video_title = f"Video_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        return {}

    def _download_video(self, url: str) -> bool:
        print_progress("Đang tải video...")
        try:
            result = subprocess.run(
                [
                    "yt-dlp",
                    "-f", f"best[height<={Config.MAX_VIDEO_HEIGHT}][ext=mp4]",
                    "-o", "temp_video.mp4",
                    url
                ],
                capture_output=True,
                text=True,
                timeout=300
            )
            if result.returncode != 0:
                print_error(f"Lỗi tải video: {result.stderr}")
                return False

            self.video_path = "temp_video.mp4"
            print_success("Video đã tải xong")
            return True
        except Exception as e:
            print_error(f"Lỗi: {e}")
            return False

    def _download_audio(self, url: str) -> bool:
        """Download audio from video"""
        print_progress("Đang tải audio...")
        try:
            audio_output = "temp_audio.m4a"
            result = subprocess.run(
                ["yt-dlp", "-f", "bestaudio[ext=m4a]", "-o", audio_output, url],
                capture_output=True,
                text=True,
                timeout=60
            )
            if result.returncode == 0 and os.path.exists(audio_output):
                self.audio_path = audio_output
                print_success("Audio đã tải xong")
                return True
            else:
                print_error("Không thể tải audio (sẽ bỏ qua transcript)")
                return False
        except Exception as e:
            print_error(f"Lỗi tải audio: {e}")
            return False

    def _transcribe_audio(self) -> bool:
        """Transcribe audio using Whisper"""
        if not self.audio_path or not os.path.exists(self.audio_path):
            print_error("Không có file audio để transcribe")
            return False

        print_progress("Đang transcribe audio...")
        try:
            with open(self.audio_path, "rb") as audio_file:
                transcript = self.client.audio.transcriptions.create(
                    model=Config.WHISPER_MODEL,
                    file=audio_file,
                    response_format="verbose_json"
                )

            self.transcript = {
                'text': transcript.text,
                'language': getattr(transcript, 'language', 'unknown'),
                'duration': getattr(transcript, 'duration', 0),
            }

            print_success(f"Transcript: {len(self.transcript['text'])} ký tự, ngôn ngữ: {self.transcript['language']}")
            return True

        except Exception as e:
            print_error(f"Lỗi transcribe: {e}")
            return False

    # ========== SCENE DETECTION ==========

    def _detect_scenes(self) -> List[Dict]:
        if not self.video_path or not os.path.exists(self.video_path):
            return []

        print_progress("Đang phát hiện scenes...")

        try:
            cap = cv2.VideoCapture(self.video_path)
            fps = cap.get(cv2.CAP_PROP_FPS)
            total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))

            if total_frames == 0:
                print_error("Không đọc được video")
                return []

            prev_frame = None
            scene_boundaries = [0]
            frame_idx = 0

            while True:
                ret, frame = cap.read()
                if not ret:
                    break

                small = cv2.resize(frame, (320, 180))
                gray = cv2.cvtColor(small, cv2.COLOR_BGR2GRAY)

                if prev_frame is not None:
                    diff = cv2.absdiff(gray, prev_frame)
                    mean_diff = np.mean(diff)

                    if mean_diff > Config.SCENE_THRESHOLD:
                        if frame_idx - scene_boundaries[-1] > Config.MIN_SCENE_LENGTH:
                            scene_boundaries.append(frame_idx)

                prev_frame = gray
                frame_idx += 1

                if frame_idx % 100 == 0:
                    progress = (frame_idx / total_frames) * 100
                    print(f"\r  Đang xử lý: {progress:.1f}%", end='', flush=True)

            print()
            scene_boundaries.append(total_frames - 1)
            cap.release()

            scenes = []
            for i in range(len(scene_boundaries) - 1):
                start = scene_boundaries[i]
                end = scene_boundaries[i + 1]
                scenes.append({
                    'scene_id': i,
                    'start_frame': start,
                    'end_frame': end,
                    'start_time': start / fps,
                    'end_time': end / fps,
                    'duration': (end - start) / fps
                })

            print_success(f"Đã phát hiện {len(scenes)} scenes")
            self.scenes = scenes
            return scenes

        except Exception as e:
            print_error(f"Lỗi phát hiện scenes: {e}")
            return []

    def _extract_first_last_frames(self) -> bool:
        """Trích xuất frame đầu và cuối"""
        if not self.scenes or not self.video_path:
            return False

        print_progress(f"Đang trích xuất frames từ {len(self.scenes)} scenes...")

        try:
            cap = cv2.VideoCapture(self.video_path)

            for scene in self.scenes:
                scene_id = scene['scene_id']
                start_frame = scene['start_frame']
                end_frame = scene['end_frame']

                # First
                cap.set(cv2.CAP_PROP_POS_FRAMES, start_frame)
                ret, frame = cap.read()
                if ret:
                    first_path = f"{Config.TEMP_FRAMES_DIR}/scene_{scene_id:04d}_FIRST.jpg"
                    cv2.imwrite(first_path, frame, [cv2.IMWRITE_JPEG_QUALITY, 95])
                    scene['frame_first'] = first_path

                # Last
                cap.set(cv2.CAP_PROP_POS_FRAMES, end_frame)
                ret, frame = cap.read()
                if ret:
                    last_path = f"{Config.TEMP_FRAMES_DIR}/scene_{scene_id:04d}_LAST.jpg"
                    cv2.imwrite(last_path, frame, [cv2.IMWRITE_JPEG_QUALITY, 95])
                    scene['frame_last'] = last_path

            cap.release()
            print_success(f"Đã trích xuất {len(self.scenes) * 2} frames")
            return True

        except Exception as e:
            print_error(f"Lỗi: {e}")
            return False

    # ========== AI ANALYSIS ==========

    def _encode_image_base64(self, image_path: str) -> str:
        with open(image_path, "rb") as f:
            return base64.standard_b64encode(f.read()).decode("utf-8")

    def _call_vision_api(self, messages: List[Dict], max_tokens: int = 800) -> Optional[str]:
        for attempt in range(Config.MAX_RETRIES):
            try:
                response = self.client.chat.completions.create(
                    model=Config.VISION_MODEL,
                    max_tokens=max_tokens,
                    messages=messages
                )
                return response.choices[0].message.content
            except Exception as e:
                if attempt < Config.MAX_RETRIES - 1:
                    print(f"  Retry {attempt + 1}...", end='', flush=True)
                    time.sleep(Config.RETRY_DELAY * (attempt + 1))
                else:
                    print_error(f"API failed: {e}")
                    return None
        return None

    def _get_frames_content(self, frame_first: str, frame_last: str) -> List[Dict]:
        """Prepare frames content for API"""
        content = []

        try:
            base64_first = self._encode_image_base64(frame_first)
            content.append({
                "type": "image_url",
                "image_url": {
                    "url": f"data:image/jpeg;base64,{base64_first}",
                    "detail": "high"
                }
            })
        except:
            pass

        try:
            base64_last = self._encode_image_base64(frame_last)
            content.append({
                "type": "image_url",
                "image_url": {
                    "url": f"data:image/jpeg;base64,{base64_last}",
                    "detail": "high"
                }
            })
        except:
            pass

        return content

    def _analyze_scene_description(self, scene: Dict) -> Optional[str]:
        """Phân tích scene để lấy miêu tả 7 mục"""
        scene_id = scene['scene_id']
        frame_first = scene.get('frame_first')
        frame_last = scene.get('frame_last')

        if not frame_first or not frame_last:
            return None

        content = [{
            "type": "text",
            "text": self.DESCRIPTION_PROMPT.format(duration=scene['duration'])
        }]
        content.extend(self._get_frames_content(frame_first, frame_last))

        description = self._call_vision_api(
            messages=[{"role": "user", "content": content}],
            max_tokens=600
        )

        return description.strip() if description else None

    def _analyze_scene_ultra_detailed(self, scene: Dict) -> Optional[Dict]:
        """Phân tích scene siêu chi tiết"""
        scene_id = scene['scene_id']
        frame_first = scene.get('frame_first')
        frame_last = scene.get('frame_last')

        if not frame_first or not frame_last:
            return None

        print_progress(f"Analyzing scene {scene_id + 1}...", scene_id + 1, len(self.scenes))

        # Step 1: Get 7-point description
        description = self._analyze_scene_description(scene)
        if description:
            scene['description'] = description
        else:
            scene['description'] = "[ERROR] Failed to generate description"

        # Step 2: Get ultra-detailed prompt
        content = [{
            "type": "text",
            "text": self.ULTRA_DETAILED_PROMPT.format(duration=scene['duration'])
        }]
        content.extend(self._get_frames_content(frame_first, frame_last))

        prompt = self._call_vision_api(
            messages=[{"role": "user", "content": content}],
            max_tokens=800
        )

        if not prompt:
            # API failed after retries - set error message
            error_msg = f"[ERROR] Failed to analyze scene {scene_id + 1} after {Config.MAX_RETRIES} retries. API may be unavailable or rate limited. Please retry later."
            print_error(f"Scene {scene_id + 1} analysis failed!")
            scene['sora_prompt'] = error_msg
            scene['analysis_failed'] = True
            return scene

        # Extract prompt
        prompt_text = prompt.strip()
        if prompt_text.startswith("PROMPT:"):
            prompt_text = prompt_text[7:].strip()

        # Validate prompt length
        word_count = len(prompt_text.split())
        if word_count < 100:
            print_error(f"Scene {scene_id + 1}: Prompt too short ({word_count} words), may be incomplete")
            scene['sora_prompt'] = f"[WARNING: Short prompt - {word_count} words only] {prompt_text}"
        else:
            scene['sora_prompt'] = prompt_text

        scene['analysis_failed'] = False
        return scene

    # ========== EXPORT TO WORD ==========

    def _export_to_word_and_txt(self) -> str:
        """Xuất ra Word (.docx) và Text, mỗi scene = 1 dòng"""
        print_progress("Đang export...")

        try:
            safe_title = "".join(c for c in self.video_title if c.isalnum() or c in (' ', '_', '-'))[:50]
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            base_name = f"{safe_title}_{timestamp}"

            output_folder = Path(Config.OUTPUT_DIR)
            output_folder.mkdir(exist_ok=True, parents=True)

            # Create separate folders for FIRST and LAST frames
            first_frames_folder = output_folder / f"{base_name}_FIRST"
            last_frames_folder = output_folder / f"{base_name}_LAST"
            first_frames_folder.mkdir(exist_ok=True)
            last_frames_folder.mkdir(exist_ok=True)

            # ========== TEXT FILE ==========
            txt_file = output_folder / f"{base_name}_PROMPTS.txt"
            with open(txt_file, 'w', encoding='utf-8') as f:
                f.write(f"VIDEO: {self.video_title}\n")
                f.write(f"URL: {self.youtube_url}\n")
                f.write(f"TOTAL SCENES: {len(self.scenes)}\n")
                f.write(f"DATE: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"{'='*70}\n\n")

                # Add transcript if available
                if self.transcript:
                    f.write(f"TRANSCRIPT:\n")
                    f.write(f"Language: {self.transcript.get('language', 'unknown')}\n")
                    f.write(f"Duration: {self.transcript.get('duration', 0):.1f}s\n")
                    f.write(f"{'-'*70}\n")
                    f.write(f"{self.transcript.get('text', '')}\n")
                    f.write(f"{'='*70}\n\n")

                for scene in self.scenes:
                    scene_num = scene['scene_id'] + 1
                    duration = scene['duration']
                    timestamp_str = f"{scene['start_time']:.1f}s - {scene['end_time']:.1f}s"
                    description = scene.get('description', 'No description')
                    prompt = scene.get('sora_prompt', 'No prompt generated')

                    # Scene header
                    f.write(f"SCENE {scene_num} ({duration:.1f}s | {timestamp_str})\n")
                    f.write(f"{'-'*70}\n\n")

                    # Description (7 mục)
                    f.write(f"MÔ TẢ CHI TIẾT:\n")
                    f.write(f"{description}\n\n")

                    # Prompt (1 dòng)
                    f.write(f"PROMPT (1 DÒNG):\n")
                    f.write(f"{prompt}\n\n")
                    f.write(f"{'='*70}\n\n")

            print_success(f"✓ Saved TXT: {txt_file.name}")

            # ========== WORD FILE (.docx) ==========
            docx_file = output_folder / f"{base_name}_PROMPTS.docx"

            print_progress("Đang tạo file Word với ảnh...")

            doc = Document()

            # Title
            title = doc.add_heading(f'Sora 2 Prompts - {self.video_title}', level=1)

            # Metadata
            meta = doc.add_paragraph()
            meta.add_run(f"Video: ").bold = True
            meta.add_run(f"{self.video_title}\n")
            meta.add_run(f"URL: ").bold = True
            meta.add_run(f"{self.youtube_url}\n")
            meta.add_run(f"Total Scenes: ").bold = True
            meta.add_run(f"{len(self.scenes)}\n")
            meta.add_run(f"Date: ").bold = True
            meta.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")

            doc.add_paragraph("_" * 50)

            # Add transcript section if available
            if self.transcript:
                doc.add_heading("Video Transcript", level=2)

                trans_meta = doc.add_paragraph()
                trans_meta.add_run("Language: ").bold = True
                trans_meta.add_run(f"{self.transcript.get('language', 'unknown')}\n")
                trans_meta.add_run("Duration: ").bold = True
                trans_meta.add_run(f"{self.transcript.get('duration', 0):.1f}s\n")

                trans_text = doc.add_paragraph()
                trans_run = trans_text.add_run(self.transcript.get('text', ''))
                trans_run.font.size = Pt(10)

                doc.add_paragraph("_" * 50)
                doc.add_paragraph()

            # Each scene = heading + images + description + prompt
            for scene in self.scenes:
                scene_num = scene['scene_id'] + 1
                duration = scene['duration']
                timestamp_str = f"{scene['start_time']:.1f}s - {scene['end_time']:.1f}s"
                description = scene.get('description', 'No description')
                prompt = scene.get('sora_prompt', 'No prompt generated')

                # Scene heading
                scene_heading = doc.add_heading(
                    f"SCENE {scene_num} ({duration:.1f}s | {timestamp_str})",
                    level=2
                )

                # Add FIRST frame image
                frame_first = scene.get('frame_first')
                if frame_first and os.path.exists(frame_first):
                    # Copy image to FIRST frames folder (numbered from 0)
                    first_dest = first_frames_folder / f"{scene['scene_id']}.jpg"
                    shutil.copy(frame_first, first_dest)

                    # Add to Word
                    p_first = doc.add_paragraph()
                    p_first.add_run("Frame đầu: ").bold = True
                    doc.add_picture(frame_first, width=Inches(4))

                # Add LAST frame image
                frame_last = scene.get('frame_last')
                if frame_last and os.path.exists(frame_last):
                    # Copy image to LAST frames folder (numbered from 0)
                    last_dest = last_frames_folder / f"{scene['scene_id']}.jpg"
                    shutil.copy(frame_last, last_dest)

                    # Add to Word
                    p_last = doc.add_paragraph()
                    p_last.add_run("Frame cuối: ").bold = True
                    doc.add_picture(frame_last, width=Inches(4))

                # Add description (7 mục)
                doc.add_paragraph()
                desc_heading = doc.add_paragraph()
                desc_heading.add_run("MÔ TẢ CHI TIẾT:").bold = True

                p_desc = doc.add_paragraph()
                run_desc = p_desc.add_run(description)
                run_desc.font.size = Pt(10)

                # Add prompt (1 line)
                doc.add_paragraph()
                prompt_heading = doc.add_paragraph()
                prompt_heading.add_run("PROMPT (1 DÒNG):").bold = True

                p_prompt = doc.add_paragraph()
                run_prompt = p_prompt.add_run(prompt)
                run_prompt.font.size = Pt(10)

                # Spacing
                doc.add_paragraph()
                doc.add_paragraph("─" * 50)
                doc.add_paragraph()

            doc.save(str(docx_file))
            print_success(f"✓ Saved WORD: {docx_file.name}")
            print_success(f"   Đường dẫn đầy đủ: {docx_file.absolute()}")

            print()
            print("="*70)
            print("📄 FILE WORD + ẢNH ĐÃ TẠO XONG!".center(70))
            print("="*70)
            print(f"\n📁 Folder chính: {output_folder.absolute()}")
            print(f"📝 File TXT: {txt_file.name}")
            print(f"📄 File WORD: {docx_file.name} (có nhúng ảnh)")
            print(f"\n📁 Folder ảnh đầu: {first_frames_folder.name}/ ({len(self.scenes)} ảnh: 0.jpg, 1.jpg, ...)")
            print(f"📁 Folder ảnh cuối: {last_frames_folder.name}/ ({len(self.scenes)} ảnh: 0.jpg, 1.jpg, ...)")
            print(f"\n✓ Tổng: {len(self.scenes)} scenes")
            print(f"✓ Mỗi scene = ảnh đầu + ảnh cuối + miêu tả 7 mục + prompt 1 dòng (250-350 words)\n")

            return str(output_folder)

        except Exception as e:
            print_error(f"Export error: {e}")
            import traceback
            traceback.print_exc()
            return ""

    # ========== MAIN PROCESS ==========

    def analyze(self, youtube_url: str) -> Optional[str]:
        """Main workflow"""
        self.youtube_url = youtube_url

        print_header("YOUTUBE TO SORA 2 - WORD EXPORT")
        print("📄 Mỗi cảnh = 1 dòng prompt siêu chi tiết")
        print("📝 Xuất ra Word (.docx) + Text (.txt)\n")

        self._get_metadata(youtube_url)

        if not self._download_video(youtube_url):
            return None

        # Download and transcribe audio
        print_header("AUDIO TRANSCRIPTION")
        if self._download_audio(youtube_url):
            self._transcribe_audio()
        else:
            print_progress("Bỏ qua transcript (không có audio)")

        if not self._detect_scenes():
            cleanup_temp()
            return None

        if not self._extract_first_last_frames():
            cleanup_temp()
            return None

        print_header(f"ANALYZING {len(self.scenes)} SCENES - ULTRA DETAILED")

        for scene in self.scenes:
            result = self._analyze_scene_ultra_detailed(scene)
            if result:
                print_success(f"Scene {scene['scene_id'] + 1} ✓")
            time.sleep(0.5)

        print_header("EXPORTING TO WORD & TEXT")
        output_path = self._export_to_word_and_txt()

        cleanup_temp()

        if output_path:
            print_header("✓ HOÀN TẤT!")
            print(f"📁 Kết quả: {output_path}/")
            print(f"📊 Tổng: {len(self.scenes)} scenes")
            print(f"📝 Files: .docx (Word) + .txt (Text)")
            print(f"🖼️  Ảnh: 2 folders riêng (FIRST + LAST)")
            print(f"📄 Mỗi scene = miêu tả 7 mục + prompt 1 dòng (250-350 words)\n")

        return output_path

# ========== CLI ==========

def main():
    print_header("YOUTUBE TO SORA 2 - WORD EXPORTER")
    print("📝 Mỗi cảnh = 1 dòng siêu chi tiết trên Word\n")

    url = None
    if len(sys.argv) > 1:
        url = sys.argv[1].strip()
        print(f"URL: {url}\n")

    if not url:
        url = input("Nhập YouTube URL: ").strip()

    if not url:
        print_error("URL không hợp lệ")
        return

    load_env_file()
    api_key = os.getenv("OPENAI_API_KEY")

    if not api_key:
        api_key = input("Nhập OpenAI API Key: ").strip()

    if not api_key:
        print_error("Cần OpenAI API Key")
        return

    try:
        analyzer = SceneBySceneWordExporter(api_key=api_key)
        result = analyzer.analyze(youtube_url=url)

        if result:
            print_success("✓ Thành công!")
        else:
            print_error("Có lỗi")

    except KeyboardInterrupt:
        print("\n\n⚠ Đã dừng")
        cleanup_temp()
    except Exception as e:
        print_error(f"Lỗi: {e}")
        import traceback
        traceback.print_exc()
        cleanup_temp()

if __name__ == "__main__":
    main()
