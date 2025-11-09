#!/usr/bin/env python3
"""
YouTube Video Analyzer to Sora 2 Prompt Generator - ADVANCED VERSION 2.0
Phân tích video YouTube chi tiết và tạo prompts chuyên nghiệp cho Sora 2

Features:
- Scene detection tự động thông minh
- Audio transcription với Whisper
- Visual composition analysis
- Camera movement detection
- Multiple prompt variants
- Intelligent caching
- Robust error handling
"""

import os
import sys
import time
import subprocess
import cv2
import base64
import numpy as np
from pathlib import Path
import json
from openai import OpenAI
from datetime import datetime
import hashlib
from typing import List, Dict, Optional, Any
import shutil

# ========== CONFIGURATION ==========

class Config:
    """Cấu hình toàn cục"""
    # Scene detection
    SCENE_THRESHOLD = 30.0
    MIN_SCENE_LENGTH = 15
    FRAMES_PER_SCENE = 3

    # Video download
    MAX_VIDEO_HEIGHT = 1080

    # API settings
    VISION_MODEL = "gpt-4o"
    TEXT_MODEL = "gpt-4o"
    WHISPER_MODEL = "whisper-1"
    MAX_RETRIES = 3
    RETRY_DELAY = 2

    # Folders
    CACHE_DIR = "cache"
    OUTPUT_DIR = "output_results"
    TEMP_FRAMES_DIR = "temp_frames"

    # Limits
    MAX_SCENES_TO_ANALYZE = 20
    MAX_SCENE_SUMMARY_LENGTH = 200


# ========== UTILITIES ==========

def load_env_file():
    """Load environment variables từ file .env"""
    env_file = Path('.env')
    if env_file.exists():
        try:
            with open(env_file, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if line and not line.startswith('#') and '=' in line:
                        key, value = line.split('=', 1)
                        key = key.strip()
                        value = value.strip()
                        # Remove quotes if present
                        if value.startswith('"') and value.endswith('"'):
                            value = value[1:-1]
                        if value.startswith("'") and value.endswith("'"):
                            value = value[1:-1]
                        os.environ[key] = value
        except Exception as e:
            print(f"⚠ Lỗi đọc file .env: {e}")


def ensure_directories():
    """Tạo các thư mục cần thiết"""
    for dir_name in [Config.CACHE_DIR, Config.OUTPUT_DIR, Config.TEMP_FRAMES_DIR]:
        Path(dir_name).mkdir(exist_ok=True)


def cleanup_temp_files():
    """Xóa tất cả files tạm"""
    temp_files = ["temp_video.mp4", "temp_audio.m4a"]
    for f in temp_files:
        if os.path.exists(f):
            try:
                os.remove(f)
            except:
                pass

    # Xóa temp frames
    if os.path.exists(Config.TEMP_FRAMES_DIR):
        try:
            shutil.rmtree(Config.TEMP_FRAMES_DIR)
            Path(Config.TEMP_FRAMES_DIR).mkdir(exist_ok=True)
        except:
            pass


def print_header(text: str):
    """In header đẹp"""
    print("\n" + "="*70)
    print(text.center(70))
    print("="*70 + "\n")


def print_section(text: str):
    """In section header"""
    print("\n" + "-"*70)
    print(text)
    print("-"*70)


def print_progress(message: str, step: Optional[int] = None, total: Optional[int] = None):
    """In progress message"""
    if step and total:
        print(f"[{step}/{total}] {message}")
    else:
        print(f"• {message}")


def print_success(message: str):
    """In success message"""
    print(f"✓ {message}")


def print_error(message: str):
    """In error message"""
    print(f"✗ {message}")


def print_warning(message: str):
    """In warning message"""
    print(f"⚠ {message}")


# ========== MAIN CLASS ==========

class YouTubeToSoraAnalyzer:
    """
    Advanced YouTube video analyzer cho Sora 2 prompt generation
    """

    def __init__(self, api_key: Optional[str] = None):
        """
        Khởi tạo analyzer

        Args:
            api_key: OpenAI API key (optional, có thể load từ env)
        """
        # Load .env first
        load_env_file()

        # Setup API key
        if api_key:
            os.environ["OPENAI_API_KEY"] = api_key

        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise ValueError(
                "OpenAI API key không tìm thấy!\n"
                "Vui lòng:\n"
                "1. Tạo file .env với: OPENAI_API_KEY=sk-your-key\n"
                "2. Hoặc set biến môi trường: export OPENAI_API_KEY=sk-your-key\n"
                "3. Hoặc truyền api_key vào constructor"
            )

        # Initialize OpenAI client
        try:
            self.client = OpenAI(api_key=api_key)
        except Exception as e:
            raise ValueError(f"Không thể khởi tạo OpenAI client: {e}")

        # Setup directories
        ensure_directories()

        # State
        self.video_path: Optional[str] = None
        self.audio_path: Optional[str] = None
        self.youtube_url: Optional[str] = None
        self.video_title: str = ""
        self.video_metadata: Dict[str, Any] = {}
        self.scenes: List[Dict] = []
        self.frames: List[str] = []

    # ========== CACHE METHODS ==========

    def _get_cache_key(self, url: str) -> str:
        """Tạo cache key từ URL"""
        return hashlib.md5(url.encode()).hexdigest()

    def _save_cache(self, key: str, data: Dict):
        """Lưu data vào cache"""
        try:
            cache_file = Path(Config.CACHE_DIR) / f"{key}.json"
            with open(cache_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print_warning(f"Không thể lưu cache: {e}")

    def _load_cache(self, key: str) -> Optional[Dict]:
        """Load data từ cache"""
        try:
            cache_file = Path(Config.CACHE_DIR) / f"{key}.json"
            if cache_file.exists():
                with open(cache_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
        except Exception as e:
            print_warning(f"Không thể đọc cache: {e}")
        return None

    # ========== VIDEO DOWNLOAD ==========

    def _get_video_metadata(self, youtube_url: str) -> Dict:
        """Lấy metadata từ YouTube"""
        print_progress("Đang lấy thông tin video...")

        try:
            result = subprocess.run(
                ["yt-dlp", "--dump-json", youtube_url],
                capture_output=True,
                text=True,
                timeout=30
            )

            if result.returncode == 0:
                metadata = json.loads(result.stdout)
                self.video_title = metadata.get('title', 'Unknown')
                self.video_metadata = {
                    'title': metadata.get('title'),
                    'duration': metadata.get('duration', 0),
                    'description': (metadata.get('description') or '')[:500],
                    'uploader': metadata.get('uploader'),
                    'width': metadata.get('width'),
                    'height': metadata.get('height'),
                    'fps': metadata.get('fps'),
                }
                print_success(f"Video: {self.video_title}")
                return self.video_metadata
        except subprocess.TimeoutExpired:
            print_error("Timeout khi lấy metadata")
        except Exception as e:
            print_error(f"Lỗi lấy metadata: {e}")

        # Fallback
        self.video_title = f"Video_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        return {}

    def _download_video(self, youtube_url: str) -> bool:
        """Download video và audio"""
        print_progress("Đang tải video...")

        try:
            video_output = "temp_video.mp4"
            audio_output = "temp_audio.m4a"

            # Download video
            result = subprocess.run(
                [
                    "yt-dlp",
                    "-f", f"best[height<={Config.MAX_VIDEO_HEIGHT}][ext=mp4]",
                    "-o", video_output,
                    youtube_url
                ],
                capture_output=True,
                text=True,
                timeout=300
            )

            if result.returncode != 0:
                print_error(f"Lỗi tải video: {result.stderr}")
                return False

            self.video_path = video_output
            print_success("Video đã tải xong")

            # Download audio (optional, for transcript)
            try:
                subprocess.run(
                    ["yt-dlp", "-f", "bestaudio[ext=m4a]", "-o", audio_output, youtube_url],
                    capture_output=True,
                    timeout=60
                )
                if os.path.exists(audio_output):
                    self.audio_path = audio_output
                    print_success("Audio đã tải xong")
            except:
                print_warning("Không thể tải audio (bỏ qua)")

            return True

        except subprocess.TimeoutExpired:
            print_error("Timeout khi tải video")
            return False
        except Exception as e:
            print_error(f"Lỗi tải video: {e}")
            return False

    # ========== SCENE DETECTION ==========

    def _detect_scenes(self) -> List[Dict]:
        """Phát hiện scenes tự động"""
        if not self.video_path or not os.path.exists(self.video_path):
            return []

        print_progress("Đang phát hiện scenes...")

        try:
            cap = cv2.VideoCapture(self.video_path)
            fps = cap.get(cv2.CAP_PROP_FPS)
            total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))

            if total_frames == 0:
                print_error("Không thể đọc video")
                return []

            prev_frame = None
            scene_boundaries = [0]
            frame_idx = 0

            while True:
                ret, frame = cap.read()
                if not ret:
                    break

                # Downsample
                small_frame = cv2.resize(frame, (320, 180))
                gray = cv2.cvtColor(small_frame, cv2.COLOR_BGR2GRAY)

                if prev_frame is not None:
                    diff = cv2.absdiff(gray, prev_frame)
                    mean_diff = np.mean(diff)

                    if mean_diff > Config.SCENE_THRESHOLD:
                        if frame_idx - scene_boundaries[-1] > Config.MIN_SCENE_LENGTH:
                            scene_boundaries.append(frame_idx)

                prev_frame = gray
                frame_idx += 1

                # Progress
                if frame_idx % 100 == 0:
                    progress = (frame_idx / total_frames) * 100
                    print(f"\r  Đang xử lý: {progress:.1f}%", end='', flush=True)

            print()  # Newline
            scene_boundaries.append(total_frames - 1)
            cap.release()

            # Create scene info
            scenes = []
            for i in range(len(scene_boundaries) - 1):
                start = scene_boundaries[i]
                end = scene_boundaries[i + 1]
                scenes.append({
                    'scene_id': i,
                    'start_frame': start,
                    'end_frame': end,
                    'start_time': start / fps,
                    'end_time': end / fps,
                    'duration': (end - start) / fps
                })

            # Limit số scenes nếu quá nhiều
            if len(scenes) > Config.MAX_SCENES_TO_ANALYZE:
                print_warning(f"Video có {len(scenes)} scenes, chỉ phân tích {Config.MAX_SCENES_TO_ANALYZE} scenes đầu")
                scenes = scenes[:Config.MAX_SCENES_TO_ANALYZE]

            self.scenes = scenes
            print_success(f"Đã phát hiện {len(scenes)} scenes")
            return scenes

        except Exception as e:
            print_error(f"Lỗi phát hiện scenes: {e}")
            return []

    def _extract_frames_from_scenes(self) -> List[str]:
        """Trích xuất key frames từ scenes"""
        if not self.scenes or not self.video_path:
            return []

        print_progress(f"Đang trích xuất {Config.FRAMES_PER_SCENE} frames từ mỗi scene...")

        try:
            cap = cv2.VideoCapture(self.video_path)
            all_frames = []

            for scene in self.scenes:
                scene_frames = []
                start = scene['start_frame']
                end = scene['end_frame']

                # Lấy frames đều
                positions = np.linspace(start, end, Config.FRAMES_PER_SCENE, dtype=int)

                for pos in positions:
                    cap.set(cv2.CAP_PROP_POS_FRAMES, pos)
                    ret, frame = cap.read()

                    if ret:
                        frame_path = f"{Config.TEMP_FRAMES_DIR}/scene_{scene['scene_id']}_frame_{len(scene_frames)}.jpg"
                        cv2.imwrite(frame_path, frame, [cv2.IMWRITE_JPEG_QUALITY, 95])
                        scene_frames.append(frame_path)

                scene['frames'] = scene_frames
                all_frames.extend(scene_frames)

            cap.release()
            self.frames = all_frames
            print_success(f"Đã trích xuất {len(all_frames)} frames")
            return all_frames

        except Exception as e:
            print_error(f"Lỗi trích xuất frames: {e}")
            return []

    # ========== AUDIO ANALYSIS ==========

    def _extract_transcript(self) -> Optional[Dict]:
        """Trích xuất transcript từ audio"""
        if not self.audio_path or not os.path.exists(self.audio_path):
            return None

        print_progress("Đang phân tích audio...")

        try:
            with open(self.audio_path, "rb") as audio_file:
                transcript = self.client.audio.transcriptions.create(
                    model=Config.WHISPER_MODEL,
                    file=audio_file,
                    response_format="verbose_json"
                )

            result = {
                'text': transcript.text,
                'language': getattr(transcript, 'language', 'unknown'),
                'duration': getattr(transcript, 'duration', 0),
            }

            print_success(f"Transcript: {len(result['text'])} ký tự")
            return result

        except Exception as e:
            print_warning(f"Không thể trích xuất transcript: {e}")
            return None

    # ========== VISUAL ANALYSIS ==========

    def _analyze_visual_composition(self, frame_path: str) -> Dict:
        """Phân tích màu sắc và composition"""
        try:
            img = cv2.imread(frame_path)
            img_rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
            pixels = img_rgb.reshape(-1, 3)

            avg_color = np.mean(pixels, axis=0)
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            brightness = float(np.mean(gray))
            contrast = float(np.std(gray))

            # Classify mood
            r, g, b = avg_color
            if brightness < 85:
                mood = "dark, moody"
            elif brightness > 170:
                mood = "bright, airy"
            elif r > g and r > b:
                mood = "warm, energetic"
            elif b > r and b > g:
                mood = "cool, calm"
            else:
                mood = "balanced, natural"

            return {
                'brightness': brightness,
                'contrast': contrast,
                'color_mood': mood
            }
        except:
            return {}

    def _encode_image_base64(self, image_path: str) -> str:
        """Encode image to base64"""
        with open(image_path, "rb") as f:
            return base64.standard_b64encode(f.read()).decode("utf-8")

    # ========== AI ANALYSIS ==========

    def _call_vision_api_with_retry(self, messages: List[Dict], max_tokens: int = 1500) -> Optional[str]:
        """Gọi Vision API với retry logic"""
        for attempt in range(Config.MAX_RETRIES):
            try:
                response = self.client.chat.completions.create(
                    model=Config.VISION_MODEL,
                    max_tokens=max_tokens,
                    messages=messages
                )
                return response.choices[0].message.content
            except Exception as e:
                if attempt < Config.MAX_RETRIES - 1:
                    print_warning(f"API call failed (attempt {attempt + 1}/{Config.MAX_RETRIES}), retrying...")
                    time.sleep(Config.RETRY_DELAY * (attempt + 1))
                else:
                    print_error(f"API call failed after {Config.MAX_RETRIES} attempts: {e}")
                    return None
        return None

    def _analyze_scene(self, scene: Dict) -> Optional[Dict]:
        """Phân tích chi tiết một scene"""
        scene_id = scene['scene_id']
        frames = scene.get('frames', [])

        if not frames:
            return None

        print_progress(f"Đang phân tích scene {scene_id + 1}...", scene_id + 1, len(self.scenes))

        # Prepare content
        content = [{
            "type": "text",
            "text": f"""Phân tích scene này (thời lượng: {scene['duration']:.1f}s):

1. HÀNH ĐỘNG: Gì đang xảy ra?
2. NHÂN VẬT/ĐỐI TƯỢNG: Ai/cái gì xuất hiện?
3. CẢM XÚC: Tâm trạng, không khí?
4. BỐI CẢNH: Địa điểm, môi trường?
5. CAMERA: Di chuyển như thế nào? (static/pan/zoom/tracking)
6. ÁNH SÁNG: Loại ánh sáng, màu sắc?
7. COMPOSITION: Cách bố cục (rule of thirds/symmetry/depth)

Trả lời ngắn gọn, chi tiết, tiếng Việt."""
        }]

        # Add frames
        for frame_path in frames:
            try:
                base64_img = self._encode_image_base64(frame_path)
                content.append({
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:image/jpeg;base64,{base64_img}",
                        "detail": "high"
                    }
                })
            except:
                continue

        # Call API
        analysis = self._call_vision_api_with_retry([{"role": "user", "content": content}])

        if not analysis:
            return None

        # Visual composition
        visual = self._analyze_visual_composition(frames[0])

        return {
            'scene_id': scene_id,
            'analysis': analysis,
            'visual_composition': visual,
            'duration': scene['duration'],
            'timestamp': f"{scene['start_time']:.1f}s - {scene['end_time']:.1f}s"
        }

    def _analyze_overall(self, scene_analyses: List[Dict], transcript: Optional[Dict]) -> Optional[str]:
        """Phân tích tổng thể video"""
        print_progress("Đang tổng hợp phân tích tổng thể...")

        # Prepare scene summaries
        scene_text = "\n\n".join([
            f"SCENE {s['scene_id'] + 1} ({s['timestamp']}):\n{s['analysis']}"
            for s in scene_analyses if s
        ])

        # Transcript
        transcript_text = ""
        if transcript:
            transcript_text = f"\n\nTRANSCRIPT:\n{transcript.get('text', '')}"

        # Metadata
        metadata_text = ""
        if self.video_metadata:
            metadata_text = f"\nThời lượng: {self.video_metadata.get('duration', 0)}s"

        prompt = f"""Dựa trên phân tích từng scene và transcript, viết tổng hợp phân tích video:

VIDEO METADATA:{metadata_text}

PHÂN TÍCH TỪNG SCENE:
{scene_text}
{transcript_text}

Hãy viết phân tích tổng thể bao gồm:
1. TÓM TẮT NỘI DUNG: Cốt truyện chính
2. PHONG CÁCH HÌNH ẢNH: Visual, màu sắc, lighting
3. KỸ THUẬT QUAY: Camera movements, transitions
4. KHÔNG KHÍ: Mood, tone, cảm xúc
5. THỂ LOẠI: Documentary/Narrative/Music Video/etc.
6. ĐẶC ĐIỂM NỔI BẬT: Điểm đặc biệt

Trả lời chi tiết, tiếng Việt."""

        return self._call_vision_api_with_retry(
            [{"role": "user", "content": prompt}],
            max_tokens=2500
        )

    def _generate_prompts(self, overall_analysis: str, scene_analyses: List[Dict]) -> Optional[str]:
        """Tạo Sora prompts"""
        print_progress("Đang tạo Sora 2 prompts...")

        # Scene summaries
        scene_text = "\n".join([
            f"Scene {s['scene_id'] + 1}: {s['analysis'][:Config.MAX_SCENE_SUMMARY_LENGTH]}..."
            for s in scene_analyses[:5] if s
        ])

        prompt = f"""Dựa trên phân tích video, tạo 3 PROMPT cho Sora 2:

PHÂN TÍCH TỔNG THỂ:
{overall_analysis}

CÁC SCENE:
{scene_text}

Tạo 3 prompts:

1. **SHORT PROMPT** (50-70 từ): Súc tích, hành động chính
2. **DETAILED PROMPT** (120-150 từ): Đầy đủ visual, camera, lighting, mood
3. **CREATIVE PROMPT** (100-130 từ): Nghệ thuật, metaphor, cinematic

YÊU CẦU:
- TẤT CẢ BẰNG TIẾNG ANH
- Rõ ràng, cụ thể, sinh động
- Mô tả camera movement, lighting, mood
- Chỉ viết prompts, không giải thích

Format:
=== SHORT PROMPT ===
[prompt]

=== DETAILED PROMPT ===
[prompt]

=== CREATIVE PROMPT ===
[prompt]"""

        return self._call_vision_api_with_retry(
            [{"role": "user", "content": prompt}],
            max_tokens=1500
        )

    # ========== EXPORT ==========

    def _save_results(self, overall: str, scenes: List[Dict], transcript: Optional[Dict], prompts: str) -> str:
        """Lưu kết quả ra files"""
        try:
            # Safe filename
            safe_title = "".join(c for c in self.video_title if c.isalnum() or c in (' ', '_', '-'))[:50]
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"{safe_title}_{timestamp}"

            # ========== TXT FILE ==========
            txt_path = f"{Config.OUTPUT_DIR}/{filename}.txt"

            scene_details = "\n\n".join([
                f"""{'='*60}
SCENE {s['scene_id'] + 1} | {s['timestamp']} | {s['duration']:.1f}s
{'='*60}

{s['analysis']}

VISUAL:
- Brightness: {s.get('visual_composition', {}).get('brightness', 0):.1f}
- Contrast: {s.get('visual_composition', {}).get('contrast', 0):.1f}
- Mood: {s.get('visual_composition', {}).get('color_mood', 'N/A')}
"""
                for s in scenes if s
            ])

            transcript_section = ""
            if transcript:
                transcript_section = f"""
{'='*60}
TRANSCRIPT
{'='*60}
Language: {transcript.get('language', 'unknown')}

{transcript.get('text', '')}
"""

            txt_content = f"""{'='*70}
YOUTUBE TO SORA 2 - ANALYSIS REPORT
{'='*70}

VIDEO: {self.video_title}
URL: {self.youtube_url}
DATE: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
Duration: {self.video_metadata.get('duration', 0)}s

{'='*70}
OVERALL ANALYSIS
{'='*70}

{overall}

{'='*70}
SCENE ANALYSIS ({len(scenes)} scenes)
{'='*70}

{scene_details}
{transcript_section}

{'='*70}
SORA 2 PROMPTS
{'='*70}

{prompts}

{'='*70}
"""

            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write(txt_content)

            print_success(f"Đã lưu TXT: {txt_path}")

            # ========== JSON FILE ==========
            json_path = f"{Config.OUTPUT_DIR}/{filename}.json"
            json_data = {
                'video_info': {
                    'title': self.video_title,
                    'url': self.youtube_url,
                    'metadata': self.video_metadata,
                    'date': datetime.now().isoformat()
                },
                'overall_analysis': overall,
                'scenes': scenes,
                'transcript': transcript,
                'sora_prompts': prompts
            }

            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(json_data, f, ensure_ascii=False, indent=2)

            print_success(f"Đã lưu JSON: {json_path}")

            # ========== DOCX FILE (optional) ==========
            try:
                from docx import Document
                from docx.enum.text import WD_ALIGN_PARAGRAPH

                doc = Document()
                title = doc.add_heading('YouTube to Sora 2 - Analysis', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                doc.add_heading('Video Information', level=1)
                table = doc.add_table(rows=3, cols=2)
                table.cell(0, 0).text = 'Title'
                table.cell(0, 1).text = self.video_title
                table.cell(1, 0).text = 'URL'
                table.cell(1, 1).text = self.youtube_url
                table.cell(2, 0).text = 'Scenes'
                table.cell(2, 1).text = str(len(scenes))

                doc.add_heading('Overall Analysis', level=1)
                doc.add_paragraph(overall)

                doc.add_heading('Scene Analysis', level=1)
                for s in scenes:
                    if s:
                        doc.add_heading(f"Scene {s['scene_id'] + 1} ({s['timestamp']})", level=2)
                        doc.add_paragraph(s['analysis'])

                if transcript:
                    doc.add_heading('Transcript', level=1)
                    doc.add_paragraph(transcript.get('text', ''))

                doc.add_heading('Sora 2 Prompts', level=1)
                doc.add_paragraph(prompts)

                docx_path = f"{Config.OUTPUT_DIR}/{filename}.docx"
                doc.save(docx_path)
                print_success(f"Đã lưu DOCX: {docx_path}")
            except ImportError:
                print_warning("Cài python-docx để export DOCX: pip install python-docx")
            except Exception as e:
                print_warning(f"Không thể tạo DOCX: {e}")

            return txt_path

        except Exception as e:
            print_error(f"Lỗi lưu file: {e}")
            return ""

    # ========== MAIN PROCESS ==========

    def analyze(self, youtube_url: str, use_cache: bool = True, analyze_audio: bool = True) -> Optional[Dict]:
        """
        Phân tích video YouTube và tạo Sora prompts

        Args:
            youtube_url: URL của video YouTube
            use_cache: Sử dụng cache nếu có
            analyze_audio: Phân tích audio/transcript

        Returns:
            Dict chứa kết quả phân tích hoặc None nếu thất bại
        """
        self.youtube_url = youtube_url

        print_header("YOUTUBE TO SORA 2 - ADVANCED ANALYZER")

        # Check cache
        cache_key = self._get_cache_key(youtube_url)
        if use_cache:
            cached = self._load_cache(cache_key)
            if cached:
                # Validate cache has complete data
                overall = cached.get('overall_analysis')
                prompts = cached.get('sora_prompts')

                if overall and prompts:
                    print_success("Tìm thấy kết quả hợp lệ trong cache!")
                    self.video_title = cached.get('video_info', {}).get('title', '')
                    print(f"Video: {self.video_title}\n")

                    print_section("PHÂN TÍCH TỔNG THỂ (từ cache)")
                    print(overall)

                    print_section("SORA 2 PROMPTS (từ cache)")
                    print(prompts)

                    return cached
                else:
                    print_warning("Cache không hợp lệ (thiếu dữ liệu), phân tích lại...")

        # Step 1: Get metadata
        self._get_video_metadata(youtube_url)

        # Step 2: Download
        if not self._download_video(youtube_url):
            return None

        # Step 3: Scene detection
        if not self._detect_scenes():
            cleanup_temp_files()
            return None

        # Step 4: Extract frames
        if not self._extract_frames_from_scenes():
            cleanup_temp_files()
            return None

        # Step 5: Transcript (optional)
        transcript = None
        if analyze_audio:
            transcript = self._extract_transcript()

        # Step 6: Analyze scenes
        print_section("PHÂN TÍCH SCENES")
        scene_analyses = []
        for scene in self.scenes:
            result = self._analyze_scene(scene)
            if result:
                scene_analyses.append(result)

        if not scene_analyses:
            print_error("Không thể phân tích scenes")
            cleanup_temp_files()
            return None

        # Step 7: Overall analysis
        print_section("PHÂN TÍCH TỔNG THỂ")
        overall = self._analyze_overall(scene_analyses, transcript)

        if overall:
            print(overall)
        else:
            cleanup_temp_files()
            return None

        # Step 8: Generate prompts
        print_section("SORA 2 PROMPTS")
        prompts = self._generate_prompts(overall, scene_analyses)

        if prompts:
            print(prompts)
        else:
            cleanup_temp_files()
            return None

        # Step 9: Save results
        print_section("LƯU KẾT QUẢ")
        self._save_results(overall, scene_analyses, transcript, prompts)

        # Save cache
        if use_cache:
            cache_data = {
                'video_info': {
                    'title': self.video_title,
                    'url': self.youtube_url,
                    'metadata': self.video_metadata
                },
                'overall_analysis': overall,
                'scene_analyses': scene_analyses,
                'transcript': transcript,
                'sora_prompts': prompts
            }
            self._save_cache(cache_key, cache_data)
            print_success("Đã lưu vào cache")

        # Cleanup
        cleanup_temp_files()

        print_header("✓ HOÀN TẤT!")
        print(f"Kết quả đã lưu trong folder: {Config.OUTPUT_DIR}/\n")

        return {
            'overall_analysis': overall,
            'scene_analyses': scene_analyses,
            'transcript': transcript,
            'sora_prompts': prompts
        }


# ========== CLI ==========

def main():
    """Main CLI interface"""
    print_header("YOUTUBE TO SORA 2 - ADVANCED ANALYZER")

    # Input URL
    youtube_url = input("Nhập YouTube URL: ").strip()
    if not youtube_url:
        print_error("URL không hợp lệ")
        return

    # Check API key
    load_env_file()
    api_key = os.getenv("OPENAI_API_KEY")

    if not api_key:
        api_key = input("Nhập OpenAI API Key: ").strip()
        if not api_key:
            print_error("Cần OpenAI API Key")
            print("\nCách lấy API key:")
            print("1. Vào https://platform.openai.com/api-keys")
            print("2. Tạo key mới")
            print("3. Lưu vào file .env hoặc nhập trực tiếp")
            return

    # Options
    print("\n--- TÙY CHỌN ---")
    use_cache = input("Sử dụng cache? (y/n, mặc định: y): ").strip().lower() != 'n'
    analyze_audio = input("Phân tích audio? (y/n, mặc định: y): ").strip().lower() != 'n'

    # Process
    try:
        analyzer = YouTubeToSoraAnalyzer(api_key=api_key)
        result = analyzer.analyze(
            youtube_url=youtube_url,
            use_cache=use_cache,
            analyze_audio=analyze_audio
        )

        if result:
            print_success("Phân tích thành công!")
        else:
            print_error("Có lỗi xảy ra")

    except Exception as e:
        print_error(f"Lỗi: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
